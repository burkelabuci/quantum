from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()
doc.add_heading('Data Table', 0)

# Define the table data
headers = ['f (GHz)', '#', 'Att (dB)', 'PSA (dBm)', 'Pant (dBm)', 'Tlaser (µs)', 'Tmw (µs)', 'LIAscale (pA)', 'Tau (ms)', 'Mode']
data = [
    [1.0, 1, 10, -50, -55, 100, 150, 10, 1.0, 'A'],
    [2.0, 2, 20, -60, -65, 200, 250, 20, 2.0, 'B'],
    [3.0, 3, 30, -70, -75, 300, 350, 30, 3.0, 'C']
]

# Add a table to the document
table = doc.add_table(rows=1, cols=len(headers))
hdr_cells = table.rows[0].cells

# Set the header cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Add data rows
for row in data:
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = str(cell_value)

# Save the document
doc.save('table.docx')
